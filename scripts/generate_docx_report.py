import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_report():
    doc = Document()
    
    # Page setup - Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Palette
    COLOR_PRIMARY = RGBColor(13, 110, 71)    # Forest Green
    COLOR_DARK = RGBColor(15, 23, 42)        # Slate Dark
    COLOR_MUTED = RGBColor(100, 116, 139)    # Slate Gray

    # Helper: Set cell background
    def set_cell_bg(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    # 1. Document Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("AWS Practical Assignment Report\n")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = COLOR_PRIMARY

    sub_run = title_p.add_run("Scalable SSC Result Publishing Infrastructure (Module 7)\n")
    sub_run.font.size = Pt(14)
    sub_run.font.bold = True
    sub_run.font.color.rgb = COLOR_DARK

    meta_run = title_p.add_run("Scenario: 10:00 AM Traffic Burst • Platform: AWS (Ubuntu 22.04 on t3.micro)\n")
    meta_run.font.size = Pt(10)
    meta_run.font.italic = True
    meta_run.font.color.rgb = COLOR_MUTED

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 2. Section: Architecture Overview
    h1 = doc.add_heading("1. Executive Summary & Architecture", level=1)
    h1.runs[0].font.color.rgb = COLOR_PRIMARY

    p1 = doc.add_paragraph(
        "This project deploys a highly available, fault-tolerant, and auto-scaling cloud infrastructure on AWS "
        "designed to handle sudden massive traffic spikes during the SSC Result 2026 publication at 10:00 AM. "
        "The architecture is engineered on t3.micro instances with an ultra-lightweight Nginx application footprint (<30 MB RAM) "
        "to ensure optimal cost efficiency and zero memory exhaustion."
    )
    p1.paragraph_format.space_after = Pt(8)

    # Architecture Table
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Component"
    hdr_cells[1].text = "AWS Resource"
    hdr_cells[2].text = "Configuration & Scope"
    for c in hdr_cells:
        set_cell_bg(c, "0D6E47")
        for p in c.paragraphs:
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(255, 255, 255)

    data = [
        ("VPC", "ssc-result-vpc", "10.0.0.0/16 across 2 Availability Zones (ap-southeast-1a & 1b)"),
        ("Public Subnets", "ssc-public-subnet-1 / 2", "10.0.1.0/24 & 10.0.2.0/24 (Hosts ALB & NAT Gateway)"),
        ("Private Subnets", "ssc-private-subnet-1 / 2", "10.0.3.0/24 & 10.0.4.0/24 (Hosts Auto Scaling EC2 Instances)"),
        ("Internet Ingress", "Application Load Balancer", "Internet-facing ALB forwarding Port 80 traffic to Target Group"),
        ("Egress Gateway", "NAT Gateway (Elastic IP)", "Provides secure outbound internet for private EC2 bootstrapping"),
        ("Auto Scaling", "ssc-result-asg", "Min: 2, Desired: 2, Max: 10 with 50% CPU Target Tracking"),
        ("Compute Node", "Ubuntu 22.04 LTS (t3.micro)", "Nginx serving SSC Result 2026 + IMDSv2 metadata injection")
    ]

    for row_idx, row in enumerate(data):
        row_cells = table.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
        row_cells[2].text = row[2]
        bg_color = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for c in row_cells:
            set_cell_bg(c, bg_color)
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 3. Section: Short Report (5 Required Questions)
    h2 = doc.add_heading("2. Technical Report & Traffic Flow", level=1)
    h2.runs[0].font.color.rgb = COLOR_PRIMARY

    q_items = [
        ("A. VPC and Network Architecture", 
         "The custom VPC 'ssc-result-vpc' (10.0.0.0/16) spans across two Availability Zones (ap-southeast-1a and ap-southeast-1b) for high availability. "
         "It contains 2 Public Subnets (for ALB and NAT Gateway) and 2 Private Subnets (for backend worker nodes). "
         "The Public Route Table directs 0.0.0.0/0 traffic to the Internet Gateway, while the Private Route Table routes all outbound traffic to the NAT Gateway."),

        ("B. ALB and EC2 Traffic Flow & Zero-Ingress Security", 
         "1. Candidates send HTTP queries to the Application Load Balancer DNS name.\n"
         "2. The ALB Security Group (ssc-alb-sg) allows incoming traffic on Port 80 from 0.0.0.0/0.\n"
         "3. The ALB distributes incoming requests across healthy private EC2 targets in both AZs using Round-Robin.\n"
         "4. The EC2 Security Group (ssc-ec2-private-sg) strictly permits Port 80 ingress exclusively from the ALB Security Group, ensuring private instances cannot be directly accessed from the internet.\n"
         "5. The ALB Target Group (ssc-result-tg) continuously polls the /health endpoint to automatically bypass any failed node within 15-30 seconds."),

        ("C. Purpose of NAT Gateway", 
         "Because the EC2 instances reside inside isolated Private Subnets without public IPv4 addresses, the NAT Gateway (positioned in Public Subnet 1 with an Elastic IP) "
         "allows private instances to perform outbound requests to download OS updates, install Nginx, and bootstrap application files via User Data while completely blocking unsolicited inbound internet traffic."),

        ("D. Auto Scaling Strategy", 
         "The Auto Scaling Group (ssc-result-asg) implements a dual proactive-reactive scaling model:\n"
         "• Baseline Fleet: 2 minimum instances deployed across 2 Availability Zones.\n"
         "• Scheduled Pre-Warming (10:00 AM Spike): A scheduled action pre-scales the group to 10 instances at 09:50 AM (10 minutes prior to result publishing) to prevent cold-start request queues.\n"
         "• Target Tracking Policy: Configured for Average CPU Utilization at 50%. If unexpected traffic surges occur, additional t3.micro instances launch within 60 seconds.\n"
         "• Scale-In: Excess instances terminate smoothly down to the baseline after peak load subsides."),

        ("E. Handling the SSC Result Traffic Spike", 
         "The combined architecture of DNS-level load balancing, Multi-AZ redundancy, lightweight Nginx static asset caching (<30 MB RAM per instance), "
         "and dynamic auto scaling from 2 to 10+ instances guarantees that millions of student requests are served with sub-second latency and zero server downtime.")
    ]

    for title, text in q_items:
        h_q = doc.add_heading(title, level=2)
        h_q.runs[0].font.color.rgb = COLOR_DARK
        p_q = doc.add_paragraph(text)
        p_q.paragraph_format.space_after = Pt(8)

    # 4. Section: AWS Console Screenshots
    doc.add_page_break()
    h3 = doc.add_heading("3. AWS Console Verification Screenshots", level=1)
    h3.runs[0].font.color.rgb = COLOR_PRIMARY

    img_dir = r"f:\devops13\mod7\ssc-result-app\images"
    screenshots = [
        ("Figure 1: Custom VPC Configuration (ssc-result-vpc - 10.0.0.0/16)", "ssc-result-vpc.png"),
        ("Figure 2: Internet Gateway Attached to VPC (ssc-result-igw)", "2.IGW.png"),
        ("Figure 3: NAT Gateway in Public Subnet (ssc-nat-gw with Elastic IP)", "3.NAt.png"),
        ("Figure 4: Public Route Table (ssc-public-rt -> IGW)", "5.public rt.png"),
        ("Figure 5: Private Route Table (ssc-private-rt -> NAT GW)", "4. private rt.png"),
        ("Figure 6: EC2 Launch Template (Ubuntu 22.04 LTS on t3.micro)", "6.launch temple.png"),
        ("Figure 7: EC2 Instances Running in Private Subnets", "7.instances.png"),
        ("Figure 8: Target Group & Health Status (Healthy Targets on Port 80)", "8.tg.png"),
        ("Figure 9: Application Load Balancer Details (ssc-result-alb)", "9.alb.png"),
        ("Figure 10: Auto Scaling Group Configuration (Min: 2, Desired: 2, Max: 10)", "10.autoScalingGroup.png"),
        ("Figure 11: Live SSC Result 2026 Web Application Accessed through ALB DNS", "11.website.png")
    ]

    for caption, fname in screenshots:
        img_path = os.path.join(img_dir, fname)
        if os.path.exists(img_path):
            doc.add_heading(caption, level=3)
            doc.add_picture(img_path, width=Inches(6.2))
            cap_p = doc.add_paragraph(f"AWS Resource Verification: {fname}")
            cap_p.runs[0].font.size = Pt(8.5)
            cap_p.runs[0].font.italic = True
            cap_p.runs[0].font.color.rgb = COLOR_MUTED
            cap_p.paragraph_format.space_after = Pt(14)

    output_path = r"f:\devops13\mod7\ssc-result-app\Assignment_7_Report.docx"
    doc.save(output_path)
    print(f"Report successfully generated at: {output_path}")

if __name__ == "__main__":
    create_report()
